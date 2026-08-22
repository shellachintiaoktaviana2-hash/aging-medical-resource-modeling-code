# -*- coding: utf-8 -*-
"""
老龄化数据与医疗资源质量层面对应关系分析
运行前请安装：pip install pandas numpy openpyxl matplotlib python-docx
将 INPUT_PATH 改为你的 Excel 文件路径后运行，会在同一文件夹输出：Excel结果表、Word描述文档和三张分析图。
"""
import sys, os, math
from pathlib import Path
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from matplotlib import font_manager
from matplotlib.patches import Patch
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
from openpyxl.drawing.image import Image as XLImage
from openpyxl.formatting.rule import ColorScaleRule
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

INPUT_PATH = Path(r"C:\Users\35442\Downloads\1老龄-质量(3).xlsx")
OUT = INPUT_PATH.parent
EXCEL_OUT = OUT / '老龄化_医疗质量对应关系分析结果.xlsx'
DOCX_OUT = OUT / '老龄化_医疗质量对应关系描述文档.docx'
FIG_GAP = OUT / '图1_Gap匹配差值排序.png'
FIG_QUAD = OUT / '图2_老龄化与医疗质量四象限.png'
FIG_COUPLE = OUT / '图3_耦合度排序.png'
TOL = 0.10

for fp in ['/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc','/usr/share/fonts/opentype/noto/NotoSerifCJK-Regular.ttc','/usr/share/fonts/truetype/arphic/uming.ttc']:
    if os.path.exists(fp):
        font_manager.fontManager.addfont(fp)
        plt.rcParams['font.sans-serif'] = [font_manager.FontProperties(fname=fp).get_name(), 'DejaVu Sans']
        break
plt.rcParams['axes.unicode_minus'] = False

raw = pd.read_excel(INPUT_PATH, sheet_name=0)
aging = raw[['地区','老龄化综合得分']].dropna().rename(columns={'老龄化综合得分':'U1_老龄化得分'})
quality = raw[['地区.1','综合得分']].dropna().rename(columns={'地区.1':'地区','综合得分':'U2_医疗质量得分'})
aging['地区'] = aging['地区'].astype(str).str.strip()
quality['地区'] = quality['地区'].astype(str).str.strip()
df = pd.merge(aging, quality, on='地区', how='inner')
df['Gap_匹配差值'] = df['U2_医疗质量得分'] - df['U1_老龄化得分']
df['Gap类型'] = np.select(
    [df['Gap_匹配差值'] < -TOL, df['Gap_匹配差值'].abs() <= TOL, df['Gap_匹配差值'] > TOL],
    ['质量资源滞后型','基本匹配型','质量资源相对富余型'], default='未分类')
den = df['U1_老龄化得分'] + df['U2_医疗质量得分']
df['C_耦合度'] = np.where(den > 0, 2*np.sqrt(df['U1_老龄化得分']*df['U2_医疗质量得分'])/den, 0)
df['耦合等级'] = pd.cut(df['C_耦合度'], [-np.inf,0.70,0.80,0.90,np.inf], labels=['低耦合','较低耦合','中等耦合','高耦合']).astype(str)
u1_mean = df['U1_老龄化得分'].mean(); u2_mean = df['U2_医疗质量得分'].mean()
u1_median = df['U1_老龄化得分'].median(); u2_median = df['U2_医疗质量得分'].median()

def quadrant(r):
    if r['U1_老龄化得分'] >= u1_mean and r['U2_医疗质量得分'] >= u2_mean: return '高老龄—高质量'
    if r['U1_老龄化得分'] >= u1_mean and r['U2_医疗质量得分'] < u2_mean: return '高老龄—低质量'
    if r['U1_老龄化得分'] < u1_mean and r['U2_医疗质量得分'] >= u2_mean: return '低老龄—高质量'
    return '低老龄—低质量'
df['四象限类型_均值分割'] = df.apply(quadrant, axis=1)
df['倾斜优先级'] = np.select([
    (df['Gap类型']=='质量资源滞后型') & (df['四象限类型_均值分割']=='高老龄—低质量'),
    (df['Gap类型']=='质量资源滞后型'),
    (df['Gap类型']=='基本匹配型'),
    (df['Gap类型']=='质量资源相对富余型')],
    ['最高优先：高老龄压力且质量资源不足','较高优先：质量资源滞后','维持优化：供需基本匹配','结构外溢/辐射：质量资源相对富余'], default='需复核')
df['老龄化排序'] = df['U1_老龄化得分'].rank(ascending=False, method='min').astype(int)
df['医疗质量排序'] = df['U2_医疗质量得分'].rank(ascending=False, method='min').astype(int)
df['Gap排序_由缺口大到富余'] = df['Gap_匹配差值'].rank(ascending=True, method='min').astype(int)
df['耦合度排序'] = df['C_耦合度'].rank(ascending=False, method='min').astype(int)
df = df[['地区','U1_老龄化得分','U2_医疗质量得分','Gap_匹配差值','Gap类型','C_耦合度','耦合等级','四象限类型_均值分割','倾斜优先级','老龄化排序','医疗质量排序','Gap排序_由缺口大到富余','耦合度排序']]
df_sorted = df.sort_values('Gap_匹配差值').reset_index(drop=True)
summary = {
    '样本数': len(df_sorted), '老龄化得分均值': u1_mean, '医疗质量得分均值': u2_mean,
    '老龄化得分中位数': u1_median, '医疗质量得分中位数': u2_median,
    'Gap均值': df_sorted['Gap_匹配差值'].mean(), 'Gap最小值': df_sorted['Gap_匹配差值'].min(), 'Gap最大值': df_sorted['Gap_匹配差值'].max(),
    '耦合度均值': df_sorted['C_耦合度'].mean(),
    'Pearson相关': df_sorted[['U1_老龄化得分','U2_医疗质量得分']].corr(method='pearson').iloc[0,1],
    'Spearman相关': df_sorted[['U1_老龄化得分','U2_医疗质量得分']].corr(method='spearman').iloc[0,1]
}
gap_counts = df_sorted['Gap类型'].value_counts().reindex(['质量资源滞后型','基本匹配型','质量资源相对富余型']).fillna(0).astype(int)
quad_counts = df_sorted['四象限类型_均值分割'].value_counts().reindex(['高老龄—高质量','高老龄—低质量','低老龄—高质量','低老龄—低质量']).fillna(0).astype(int)
couple_counts = df_sorted['耦合等级'].value_counts().reindex(['高耦合','中等耦合','较低耦合','低耦合']).fillna(0).astype(int)

# Figures
colors_map = {'质量资源滞后型':'#d55e00','基本匹配型':'#777777','质量资源相对富余型':'#0072b2'}
fig, ax = plt.subplots(figsize=(11.8,6.5), dpi=220)
ax.bar(df_sorted['地区'], df_sorted['Gap_匹配差值'], color=[colors_map[x] for x in df_sorted['Gap类型']])
ax.axhline(0, color='black', lw=.9); ax.axhline(TOL, color='#999999', lw=.8, ls='--'); ax.axhline(-TOL, color='#999999', lw=.8, ls='--')
ax.set_title('各地区医疗质量相对老龄化压力的匹配差值（Gap = U2 - U1）', fontsize=14, pad=14)
ax.set_ylabel('Gap匹配差值'); ax.set_xlabel('地区（按Gap由低到高排序）'); ax.tick_params(axis='x', rotation=70); ax.grid(axis='y', alpha=.25)
ax.legend(handles=[Patch(facecolor=colors_map[k], label=k) for k in colors_map], loc='upper left', frameon=False, ncols=3)
fig.tight_layout(); fig.savefig(FIG_GAP, bbox_inches='tight'); plt.close(fig)
quad_color = {'高老龄—高质量':'#009e73','高老龄—低质量':'#d55e00','低老龄—高质量':'#0072b2','低老龄—低质量':'#999999'}
fig, ax = plt.subplots(figsize=(9.5,7.3), dpi=220)
for q, sub in df_sorted.groupby('四象限类型_均值分割'):
    ax.scatter(sub['U1_老龄化得分'], sub['U2_医疗质量得分'], s=64, label=q, color=quad_color[q], edgecolor='white', linewidth=.7)
for _, r in df_sorted.iterrows(): ax.text(r['U1_老龄化得分']+0.008, r['U2_医疗质量得分']+0.006, r['地区'], fontsize=8.2)
ax.axvline(u1_mean, color='black', lw=.9, ls='--'); ax.axhline(u2_mean, color='black', lw=.9, ls='--')
ax.set_title('老龄化得分与医疗质量得分四象限分布', fontsize=14, pad=14)
ax.set_xlabel(f'U1 老龄化综合得分（均值={u1_mean:.4f}）'); ax.set_ylabel(f'U2 医疗质量综合得分（均值={u2_mean:.4f}）')
ax.set_xlim(-.03, 1.05); ax.set_ylim(0, .82); ax.legend(loc='lower right', fontsize=9); ax.grid(alpha=.25)
fig.tight_layout(); fig.savefig(FIG_QUAD, bbox_inches='tight'); plt.close(fig)
couple_sorted = df_sorted.sort_values('C_耦合度', ascending=False)
fig, ax = plt.subplots(figsize=(11.8,6.2), dpi=220)
bar_cols = ['#009e73' if x=='高耦合' else '#e69f00' if x=='中等耦合' else '#999999' if x=='较低耦合' else '#d55e00' for x in couple_sorted['耦合等级']]
ax.bar(couple_sorted['地区'], couple_sorted['C_耦合度'], color=bar_cols)
ax.axhline(.9, color='#999999', ls='--', lw=.8); ax.axhline(.8, color='#bbbbbb', ls='--', lw=.8)
ax.set_title('老龄化—医疗质量耦合度排序', fontsize=14, pad=14); ax.set_ylabel('耦合度 C'); ax.set_xlabel('地区（按耦合度由高到低排序）')
ax.set_ylim(0,1.08); ax.tick_params(axis='x', rotation=70); ax.grid(axis='y', alpha=.25)
fig.tight_layout(); fig.savefig(FIG_COUPLE, bbox_inches='tight'); plt.close(fig)

# Excel
wb = Workbook(); ws = wb.active; ws.title='最终结果表'
ws.merge_cells('A1:M1'); ws['A1']='老龄化得分—医疗质量层面对应关系分析结果'; ws['A1'].font=Font(name='Microsoft YaHei',size=16,bold=True,color='1F2937'); ws['A1'].alignment=Alignment(horizontal='center',vertical='center'); ws.row_dimensions[1].height=30
ws.append([]); ws.append(list(df_sorted.columns))
for c in ws[3]: c.font=Font(name='Microsoft YaHei',bold=True,color='1F2937'); c.fill=PatternFill('solid',fgColor='E8F0FE'); c.alignment=Alignment(horizontal='center',vertical='center',wrap_text=True); c.border=Border(bottom=Side(style='thin',color='AAB7C4'))
for row in df_sorted.itertuples(index=False): ws.append(list(row))
for row in ws.iter_rows(min_row=4,max_row=ws.max_row,min_col=1,max_col=ws.max_column):
    for c in row: c.font=Font(name='Microsoft YaHei',size=10,color='111827'); c.alignment=Alignment(horizontal='center',vertical='center',wrap_text=True); c.border=Border(bottom=Side(style='thin',color='E5E7EB'))
for col in ['B','C','D','F']:
    for c in ws[col][3:]: c.number_format='0.0000'
fill_gap={'质量资源滞后型':'FCE4D6','基本匹配型':'F2F2F2','质量资源相对富余型':'DDEBF7'}; fill_quad={'高老龄—低质量':'F8CBAD','高老龄—高质量':'E2F0D9','低老龄—高质量':'DDEBF7','低老龄—低质量':'E7E6E6'}
for r in range(4, ws.max_row+1):
    ws[f'E{r}'].fill=PatternFill('solid',fgColor=fill_gap.get(ws[f'E{r}'].value,'FFFFFF')); ws[f'H{r}'].fill=PatternFill('solid',fgColor=fill_quad.get(ws[f'H{r}'].value,'FFFFFF'))
    val=str(ws[f'I{r}'].value)
    ws[f'I{r}'].fill=PatternFill('solid',fgColor='F4CCCC' if val.startswith('最高优先') else 'FCE4D6' if val.startswith('较高优先') else 'E7E6E6' if val.startswith('维持优化') else 'D9EAF7')
ws.freeze_panes='A4'; ws.auto_filter.ref=f'A3:M{ws.max_row}'
for col,w in {'A':10,'B':15,'C':16,'D':14,'E':18,'F':11,'G':11,'H':18,'I':30,'J':10,'K':12,'L':17,'M':11}.items(): ws.column_dimensions[col].width=w
for r in range(4, ws.max_row+1): ws.row_dimensions[r].height=24
ws.conditional_formatting.add(f'D4:D{ws.max_row}', ColorScaleRule(start_type='min',start_color='F4CCCC',mid_type='num',mid_value=0,mid_color='FFFFFF',end_type='max',end_color='CFE2F3'))
ws.conditional_formatting.add(f'F4:F{ws.max_row}', ColorScaleRule(start_type='min',start_color='F4CCCC',end_type='max',end_color='D9EAD3'))
ws2=wb.create_sheet('方法与汇总'); ws2.merge_cells('A1:C1'); ws2['A1']='分析方法、公式与总体结论'; ws2['A1'].font=Font(name='Microsoft YaHei',size=16,bold=True,color='1F2937'); ws2['A1'].alignment=Alignment(horizontal='center')
for row in [['步骤','公式/规则','说明'],['1. 数据匹配','按“地区”合并 U1 与 U2','U1为老龄化综合得分；U2为医疗质量综合得分。'],['2. Gap差值','Gap_i = U2_i - U1_i',f'Gap<-{TOL:.2f}为质量资源滞后型；|Gap|≤{TOL:.2f}为基本匹配型；Gap>{TOL:.2f}为相对富余型。'],['3. 耦合度','C_i = 2√(U1_i×U2_i)/(U1_i+U2_i)','C越接近1，说明两系统数值同步性越强；但不直接代表资源是否充足。'],['4. 四象限',f'U1均值={u1_mean:.4f}；U2均值={u2_mean:.4f}','以均值划分高/低，形成四类。']]: ws2.append(row)
ws2.append([]); ws2.append(['统计项','数值'])
for k,v in summary.items(): ws2.append([k,float(v) if isinstance(v,(float,np.floating,int)) else v])
for row in ws2.iter_rows(min_row=2,max_row=ws2.max_row,min_col=1,max_col=3):
    for c in row: c.font=Font(name='Microsoft YaHei',size=10,bold=(c.row in [2,8]),color='111827'); c.alignment=Alignment(vertical='center',wrap_text=True); c.border=Border(bottom=Side(style='thin',color='E5E7EB')); c.fill=PatternFill('solid',fgColor='E8F0FE' if c.row in [2,8] else 'FFFFFF')
for col,w in {'A':23,'B':30,'C':60}.items(): ws2.column_dimensions[col].width=w
ws3=wb.create_sheet('分类汇总'); ws3['A1']='分类汇总'; ws3['A1'].font=Font(name='Microsoft YaHei',size=15,bold=True)
row=3
for title, series in [('Gap类型分布',gap_counts),('四象限分布',quad_counts),('耦合等级分布',couple_counts)]:
    ws3[f'A{row}']=title; ws3[f'A{row}'].font=Font(name='Microsoft YaHei',size=12,bold=True); row+=1
    ws3[f'A{row}']='类别'; ws3[f'B{row}']='数量'; ws3[f'A{row}'].fill=ws3[f'B{row}'].fill=PatternFill('solid',fgColor='E8F0FE'); row+=1
    for idx,val in series.items(): ws3[f'A{row}']=idx; ws3[f'B{row}']=int(val); row+=1
    row+=2
for r in ws3.iter_rows():
    for c in r: c.font=Font(name='Microsoft YaHei',size=10,bold=c.font.bold); c.alignment=Alignment(vertical='center');
ws3.column_dimensions['A'].width=30; ws3.column_dimensions['B'].width=10
ws4=wb.create_sheet('图形展示'); ws4['A1']='图形展示：Gap排序、四象限与耦合度'; ws4['A1'].font=Font(name='Microsoft YaHei',size=15,bold=True)
for img_path, cell, width, height in [(FIG_GAP,'B3',900,500),(FIG_QUAD,'B32',720,550),(FIG_COUPLE,'B66',900,480)]:
    img=XLImage(str(img_path)); img.width=width; img.height=height; ws4.add_image(img, cell)
ws5=wb.create_sheet('原始匹配数据'); ws5.append(['地区','U1_老龄化得分','U2_医疗质量得分'])
for r in df[['地区','U1_老龄化得分','U2_医疗质量得分']].sort_values('地区').itertuples(index=False): ws5.append(list(r))
for row in ws5.iter_rows():
    for c in row: c.font=Font(name='Microsoft YaHei',size=10,bold=(c.row==1)); c.alignment=Alignment(horizontal='center'); c.border=Border(bottom=Side(style='thin',color='E5E7EB'))
for c in ws5[1]: c.fill=PatternFill('solid',fgColor='E8F0FE')
for col,w in {'A':18,'B':18,'C':18}.items(): ws5.column_dimensions[col].width=w
for col in ['B','C']:
    for c in ws5[col][1:]: c.number_format='0.0000'
wb.save(EXCEL_OUT)

# Word

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)
def cell_text(cell, text, bold=False):
    cell.text=''; p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(str(text)); r.font.name='Microsoft YaHei'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei'); r.font.size=Pt(8.5); r.bold=bold
def add_df_table(doc, d, title=None):
    if title: doc.add_heading(title, level=3)
    table=doc.add_table(rows=1, cols=len(d.columns)); table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.style='Table Grid'
    for i,col in enumerate(d.columns): cell_text(table.rows[0].cells[i], col, True); shade(table.rows[0].cells[i], 'E8F0FE')
    for _,row in d.iterrows():
        cells=table.add_row().cells
        for j,val in enumerate(row): cell_text(cells[j], f'{val:.4f}' if isinstance(val,(float,np.floating)) else val)
    for rr in table.rows:
        for cc in rr.cells: cc.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()

doc=Document();
sec=doc.sections[0]; sec.top_margin=Inches(.75); sec.bottom_margin=Inches(.75); sec.left_margin=Inches(.78); sec.right_margin=Inches(.78)
for st in ['Normal','Heading 1','Heading 2','Heading 3']:
    doc.styles[st].font.name='Microsoft YaHei'; doc.styles[st]._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei')
doc.styles['Normal'].font.size=Pt(10.5)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('老龄化数据与医疗资源质量层面对应关系分析'); r.bold=True; r.font.size=Pt(18); r.font.name='Microsoft YaHei'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('基于 Gap 匹配差值、耦合度与四象限分类的综合分析'); r.font.size=Pt(11); r.font.color.rgb=RGBColor(75,85,99); r.font.name='Microsoft YaHei'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei')
doc.add_paragraph('本报告使用上传 Excel 中的两组综合得分进行匹配分析：U1 表示老龄化综合得分，U2 表示医疗资源质量层面综合得分。样本共覆盖 31 个地区，分析目标是识别医疗质量资源相对于老龄化压力的滞后、匹配或相对富余状态，并进一步判断各地区在供需关系中的象限位置。')
doc.add_heading('一、数据处理与指标定义', level=1)
for txt in ['1. 数据读取：从工作表“综合得分排名”中提取左侧“地区—老龄化综合得分”和右侧“地区—综合得分”两张表。','2. 地区匹配：以“地区”为主键进行内连接，确保每个地区同时具有老龄化得分 U1 和医疗质量得分 U2。','3. 得分方向：U1 越高表示老龄化压力越高；U2 越高表示医疗资源质量层面综合水平越高。']: doc.add_paragraph(txt)
doc.add_heading('二、第一层：主分析，使用匹配差值 Gap', level=1)
doc.add_paragraph('计算公式：Gap_i = U2_i − U1_i。')
doc.add_paragraph(f'判定规则采用 ±{TOL:.2f} 作为实际匹配容差：Gap < -{TOL:.2f} 为“质量资源滞后型”；|Gap| ≤ {TOL:.2f} 为“基本匹配型”；Gap > {TOL:.2f} 为“质量资源相对富余型”。该阈值表示在 0—1 综合得分体系中允许约 10% 的轻微波动。')
doc.add_paragraph(f'结果显示：31 个地区中，质量资源滞后型 {int(gap_counts["质量资源滞后型"])} 个，基本匹配型 {int(gap_counts["基本匹配型"])} 个，质量资源相对富余型 {int(gap_counts["质量资源相对富余型"])} 个。Gap 均值为 {summary["Gap均值"]:.4f}，整体上医疗质量层面相对于老龄化压力略偏滞后。')
add_df_table(doc, df_sorted[['地区','U1_老龄化得分','U2_医疗质量得分','Gap_匹配差值','Gap类型','倾斜优先级']].head(10), '表1 质量资源滞后程度前10位地区')
doc.add_picture(str(FIG_GAP), width=Inches(6.5)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
doc.add_heading('三、第二层：辅助分析，使用耦合度 C', level=1)
doc.add_paragraph('计算公式：C_i = 2√(U1_i × U2_i) / (U1_i + U2_i)。')
doc.add_paragraph('耦合度用于描述老龄化系统与医疗质量系统的同步程度，C 越接近 1，说明两者数值越接近；但耦合度不直接代表资源是否充足，因此需要与 Gap 共同解释。')
doc.add_paragraph(f'本次结果中，平均耦合度为 {summary["耦合度均值"]:.4f}；高耦合地区 {int(couple_counts["高耦合"])} 个，中等耦合地区 {int(couple_counts["中等耦合"])} 个，低耦合地区 {int(couple_counts["低耦合"])} 个。低耦合主要出现在 U1 与 U2 差异极大的地区。')
add_df_table(doc, df_sorted.sort_values('C_耦合度', ascending=False)[['地区','U1_老龄化得分','U2_医疗质量得分','Gap_匹配差值','C_耦合度','耦合等级']].head(10), '表2 耦合度前10位地区')
doc.add_picture(str(FIG_COUPLE), width=Inches(6.5)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
doc.add_heading('四、第三层：图形展示，使用四象限分类', level=1)
doc.add_paragraph(f'四象限以样本均值作为高低划分阈值：老龄化得分均值为 {u1_mean:.4f}，医疗质量得分均值为 {u2_mean:.4f}。横轴为 U1，纵轴为 U2。')
doc.add_paragraph(f'四象限结果为：高老龄—高质量 {int(quad_counts["高老龄—高质量"])} 个，高老龄—低质量 {int(quad_counts["高老龄—低质量"])} 个，低老龄—高质量 {int(quad_counts["低老龄—高质量"])} 个，低老龄—低质量 {int(quad_counts["低老龄—低质量"])} 个。')
quad_table=pd.DataFrame({'象限类型':quad_counts.index,'地区数量':quad_counts.values,'典型含义':['老龄化压力较高，但医疗质量基础相对较强；重点在结构效率和服务可及性。','老龄化压力较高且医疗质量低于均值；属于优先补短板区域。','老龄化压力低于均值、医疗质量高于均值；可发挥区域辐射或资源外溢作用。','老龄化压力和医疗质量均低于均值；需关注潜在老龄化上升后的承载能力。']})
add_df_table(doc, quad_table, '表3 四象限类型及含义')
doc.add_picture(str(FIG_QUAD), width=Inches(6.3)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
doc.add_heading('五、重点结论与解释', level=1)
hp_names='、'.join(df_sorted[df_sorted['倾斜优先级'].str.startswith('最高优先')]['地区'].tolist())
surplus_names='、'.join(df_sorted[df_sorted['Gap类型']=='质量资源相对富余型']['地区'].tolist())
match_names='、'.join(df_sorted[df_sorted['Gap类型']=='基本匹配型']['地区'].tolist())
for txt in [f'1. 需要重点倾斜的地区为：{hp_names}。这些地区同时满足“高老龄—低质量”或明显质量资源滞后，说明老龄化压力较高，但医疗质量层面的综合支撑不足。','2. 辽宁、吉林、河北、重庆、江苏、山东等地区的 Gap 为明显负值，说明医疗质量得分相对于老龄化得分存在较大落差，是优先补齐质量型医疗资源、提升服务能力和资源等级结构的重点对象。',f'3. 基本匹配地区为：{match_names}。这些地区的医疗质量得分与老龄化得分差距较小，后续应以结构优化和稳定供给为主。',f'4. 质量资源相对富余地区为：{surplus_names}。这些地区的 U2 高于 U1，短期内质量资源相对充足，但仍需结合人口流动、服务半径和区域协同能力进一步判断是否可承担辐射功能。',f'5. 补充相关性结果显示，Pearson 相关系数为 {summary["Pearson相关"]:.4f}，Spearman 相关系数为 {summary["Spearman相关"]:.4f}，表明当前老龄化水平与医疗质量得分之间并不存在明显单调同步关系。']: doc.add_paragraph(txt)
doc.add_heading('六、方法适用性与注意事项', level=1)
for txt in ['1. Gap 是识别“谁该优先倾斜”的主依据，因为它直接给出了医疗质量相对老龄化压力的差额方向。','2. 耦合度是辅助依据，适合描述两组得分是否同步，但高耦合可能同时出现在“双高”或“双低”地区，不能单独作为资源充足性判断。','3. 四象限适合做政策解释和图形展示，能够把地区归入“高老龄—高质量”“高老龄—低质量”“低老龄—高质量”“低老龄—低质量”四类。','4. 本报告使用均值作为四象限阈值；若比赛论文中希望弱化极端值影响，也可以在附录中改用中位数阈值进行稳健性检验。']: doc.add_paragraph(txt)
doc.add_heading('七、最终结果表说明', level=1)
doc.add_paragraph('完整结果已输出至 Excel 文件“老龄化_医疗质量对应关系分析结果.xlsx”。其中“最终结果表”包含地区、U1、U2、Gap、Gap类型、耦合度、耦合等级、四象限类型、倾斜优先级及排序字段；“方法与汇总”保存公式、阈值和总体统计；“图形展示”保存三张核心图。')
doc.save(DOCX_OUT)
