# -*- coding: utf-8 -*-
"""
老龄化数据与医疗资源空间层面对应关系分析
依赖安装：pip install pandas openpyxl matplotlib python-docx

使用方法：
1. 将本脚本与“1老龄-空间(3).xlsx”放在同一文件夹；或修改 INPUT_FILE 为你的 Excel 绝对路径。
2. 运行后会在同一文件夹输出：
   - 老龄化_医疗空间层面对应关系分析结果.xlsx
   - 老龄化_医疗空间四象限图.png
   - 老龄化_医疗空间层面对应关系分析说明.docx
"""

from pathlib import Path
import math
from collections import Counter, defaultdict

import pandas as pd
import matplotlib.pyplot as plt
from matplotlib import font_manager
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import load_workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.formatting.rule import ColorScaleRule, DataBarRule
from openpyxl.chart import ScatterChart, Series, Reference

# =========================
# 1. 参数区
# =========================
SCRIPT_DIR = Path(__file__).resolve().parent
INPUT_FILE = SCRIPT_DIR / "1老龄-空间(3).xlsx"  # 如需绝对路径，可改为 Path(r"C:\\Users\\35442\\Downloads\\1老龄-空间(3).xlsx")
OUTPUT_EXCEL = SCRIPT_DIR / "老龄化_医疗空间层面对应关系分析结果.xlsx"
OUTPUT_PNG = SCRIPT_DIR / "老龄化_医疗空间四象限图.png"
OUTPUT_DOCX = SCRIPT_DIR / "老龄化_医疗空间层面对应关系分析说明.docx"
GAP_THRESHOLD = 0.05


# =========================
# 2. 数据读取与整理
# =========================
def read_score_data(input_file: Path):
    """读取上传结果表。默认 A-B 列为老龄化结果，F-G 列为空间维度医疗结果。"""
    raw = pd.read_excel(input_file, sheet_name=0, header=0)

    aging = raw.iloc[:, [0, 1]].dropna()
    aging.columns = ["地区", "U1老龄化综合得分"]

    spatial = raw.iloc[:, [5, 6]].dropna()
    spatial.columns = ["地区", "U2医疗空间综合得分"]

    aging["U1老龄化综合得分"] = pd.to_numeric(aging["U1老龄化综合得分"], errors="coerce")
    spatial["U2医疗空间综合得分"] = pd.to_numeric(spatial["U2医疗空间综合得分"], errors="coerce")
    aging = aging.dropna()
    spatial = spatial.dropna()

    df = pd.merge(aging, spatial, on="地区", how="inner")
    if df.empty:
        raise ValueError("未能按‘地区’匹配到有效数据，请检查原始表的 A-B 列和 F-G 列。")
    return df


# =========================
# 3. 指标计算与分类
# =========================
def gap_type(gap: float, threshold: float = GAP_THRESHOLD) -> str:
    if gap < -threshold:
        return "空间资源滞后型"
    if gap > threshold:
        return "空间资源相对富余型"
    return "基本匹配型"


def coupling_grade(c: float) -> str:
    if c < 0.3:
        return "低耦合"
    if c < 0.5:
        return "中低耦合"
    if c < 0.8:
        return "中等耦合"
    return "高耦合"


def add_model_results(df: pd.DataFrame) -> pd.DataFrame:
    """按图片方法计算 Gap、耦合度 C 和四象限分类。"""
    result = df.copy()
    result["U1排名"] = result["U1老龄化综合得分"].rank(ascending=False, method="min").astype(int)
    result["U2排名"] = result["U2医疗空间综合得分"].rank(ascending=False, method="min").astype(int)

    # 第一层：匹配差值
    result["Gap=U2-U1"] = result["U2医疗空间综合得分"] - result["U1老龄化综合得分"]
    result["匹配类型"] = result["Gap=U2-U1"].apply(gap_type)

    # 第二层：耦合度
    def coupling(row):
        u1 = row["U1老龄化综合得分"]
        u2 = row["U2医疗空间综合得分"]
        return 0 if (u1 + u2) == 0 else 2 * math.sqrt(max(u1, 0) * max(u2, 0)) / (u1 + u2)

    result["耦合度C"] = result.apply(coupling, axis=1)
    result["耦合等级"] = result["耦合度C"].apply(coupling_grade)

    # 第三层：四象限
    u1_mean = result["U1老龄化综合得分"].mean()
    u2_mean = result["U2医疗空间综合得分"].mean()

    def quadrant(row):
        u1 = row["U1老龄化综合得分"]
        u2 = row["U2医疗空间综合得分"]
        if u1 >= u1_mean and u2 >= u2_mean:
            return "高老龄—高空间资源"
        if u1 >= u1_mean and u2 < u2_mean:
            return "高老龄—低空间资源"
        if u1 < u1_mean and u2 >= u2_mean:
            return "低老龄—高空间资源"
        return "低老龄—低空间资源"

    result["四象限类型"] = result.apply(quadrant, axis=1)

    def conclusion(row):
        if row["匹配类型"] == "空间资源滞后型" and row["四象限类型"] == "高老龄—低空间资源":
            return "重点补短板：老龄化压力较高，空间资源供给偏弱。"
        if row["匹配类型"] == "空间资源滞后型" and row["四象限类型"] == "高老龄—高空间资源":
            return "总量相对较高但仍低于老龄化压力，应优化空间覆盖与可及性。"
        if row["匹配类型"] == "空间资源相对富余型":
            return "空间资源相对富余，需关注利用效率与周边辐射服务。"
        if row["匹配类型"] == "基本匹配型":
            return "供需差距较小，属于当前样本内基本匹配。"
        return "需结合人口密度、地理面积和服务半径进一步解释。"

    result["识别结论"] = result.apply(conclusion, axis=1)
    result = result.sort_values("Gap=U2-U1", ascending=True).reset_index(drop=True)
    return result


# =========================
# 4. 输出图、Excel、Word
# =========================
def make_quadrant_plot(result: pd.DataFrame, output_png: Path):
    # 尽量设置中文字体；若本机没有这些字体，图仍会生成，但中文可能显示为方框。
    font_paths = [
        Path(r"C:\Windows\Fonts\msyh.ttc"),      # Windows：微软雅黑
        Path(r"C:\Windows\Fonts\simhei.ttf"),    # Windows：黑体
        Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),  # Linux
        Path("/System/Library/Fonts/PingFang.ttc"),    # macOS
    ]
    for fp in font_paths:
        if fp.exists():
            font_manager.fontManager.addfont(str(fp))
            prop = font_manager.FontProperties(fname=str(fp))
            plt.rcParams["font.family"] = prop.get_name()
            break
    else:
        font_candidates = ["Microsoft YaHei", "SimHei", "Noto Sans CJK SC", "Source Han Sans SC"]
        available_fonts = {f.name for f in font_manager.fontManager.ttflist}
        for f in font_candidates:
            if f in available_fonts:
                plt.rcParams["font.sans-serif"] = [f]
                break
    plt.rcParams["axes.unicode_minus"] = False

    u1_mean = result["U1老龄化综合得分"].mean()
    u2_mean = result["U2医疗空间综合得分"].mean()
    color_map = {
        "高老龄—高空间资源": "#1f77b4",
        "高老龄—低空间资源": "#d62728",
        "低老龄—高空间资源": "#2ca02c",
        "低老龄—低空间资源": "#7f7f7f",
    }

    plt.figure(figsize=(10, 7), dpi=180)
    for q, sub in result.groupby("四象限类型"):
        plt.scatter(sub["U1老龄化综合得分"], sub["U2医疗空间综合得分"],
                    s=70, alpha=0.85, label=f"{q}（{len(sub)}）", color=color_map.get(q, "#333333"))

    label_names = set(result.nsmallest(5, "Gap=U2-U1")["地区"]) | set(result.nlargest(4, "Gap=U2-U1")["地区"])
    label_names |= {"上海", "北京", "河南", "江苏"}
    for _, row in result.iterrows():
        if row["地区"] in label_names:
            plt.annotate(row["地区"], (row["U1老龄化综合得分"], row["U2医疗空间综合得分"]),
                         textcoords="offset points", xytext=(5, 4), fontsize=8)

    plt.axvline(u1_mean, color="#444444", linestyle="--", linewidth=1.2)
    plt.axhline(u2_mean, color="#444444", linestyle="--", linewidth=1.2)
    plt.text(u1_mean + 0.01, 0.96, f"U1均值={u1_mean:.3f}", fontsize=9)
    plt.text(0.02, u2_mean + 0.02, f"U2均值={u2_mean:.3f}", fontsize=9)
    plt.xlabel("U1 老龄化综合得分")
    plt.ylabel("U2 医疗空间维度综合得分")
    plt.title("老龄化数据与医疗空间资源的四象限对应关系")
    plt.xlim(-0.03, 1.05)
    plt.ylim(0, 1.05)
    plt.grid(True, linestyle=":", alpha=0.35)
    plt.legend(loc="upper left", frameon=True)
    plt.tight_layout()
    plt.savefig(output_png, bbox_inches="tight")
    plt.close()


def write_excel(result: pd.DataFrame, output_excel: Path):
    u1_mean = result["U1老龄化综合得分"].mean()
    u2_mean = result["U2医疗空间综合得分"].mean()
    gap_mean = result["Gap=U2-U1"].mean()
    c_mean = result["耦合度C"].mean()

    gap_counts = result["匹配类型"].value_counts()
    quad_counts = result["四象限类型"].value_counts()
    c_counts = result["耦合等级"].value_counts()
    quad_regions = result.groupby("四象限类型")["地区"].apply(lambda x: "、".join(x)).to_dict()

    method_df = pd.DataFrame({
        "项目": ["研究对象", "第一层：匹配差值", "Gap解释", "第二层：耦合度", "耦合等级", "第三层：四象限"],
        "说明/公式": [
            "31个省级地区，按地区名称匹配老龄化综合得分U1与医疗空间维度得分U2。",
            "Gap_i = U_2i - U_1i",
            "Gap<0：空间资源滞后；Gap≈0：基本匹配；Gap>0：空间资源相对富余。",
            "C_i = 2√(U_1i U_2i)/(U_1i + U_2i)",
            "C<0.3低耦合；0.3≤C<0.5中低耦合；0.5≤C<0.8中等耦合；C≥0.8高耦合。",
            "横轴U1，纵轴U2，以样本均值作为分界。",
        ],
        "本次设置": [
            str(INPUT_FILE),
            f"|Gap|≤{GAP_THRESHOLD} 判为基本匹配",
            "用于识别资源应优先倾斜的主体",
            "C越接近1，两个系统数值越同步",
            "耦合度只表示同步性，不等同于资源充足",
            f"U1均值={u1_mean:.6f}；U2均值={u2_mean:.6f}",
        ]
    })

    summary_rows = [
        ["U1老龄化均值", u1_mean],
        ["U2医疗空间均值", u2_mean],
        ["Gap均值", gap_mean],
        ["耦合度C均值", c_mean],
        ["空间资源滞后型数量", gap_counts.get("空间资源滞后型", 0)],
        ["基本匹配型数量", gap_counts.get("基本匹配型", 0)],
        ["空间资源相对富余型数量", gap_counts.get("空间资源相对富余型", 0)],
        ["高老龄—高空间资源", f"{quad_counts.get('高老龄—高空间资源', 0)}：{quad_regions.get('高老龄—高空间资源', '')}"],
        ["高老龄—低空间资源", f"{quad_counts.get('高老龄—低空间资源', 0)}：{quad_regions.get('高老龄—低空间资源', '')}"],
        ["低老龄—高空间资源", f"{quad_counts.get('低老龄—高空间资源', 0)}：{quad_regions.get('低老龄—高空间资源', '')}"],
        ["低老龄—低空间资源", f"{quad_counts.get('低老龄—低空间资源', 0)}：{quad_regions.get('低老龄—低空间资源', '')}"],
    ]
    summary_df = pd.DataFrame(summary_rows, columns=["项目", "结果"])

    chart_df = result[["地区", "U1老龄化综合得分", "U2医疗空间综合得分", "四象限类型", "Gap=U2-U1", "耦合度C"]].sort_values("地区")

    with pd.ExcelWriter(output_excel, engine="openpyxl") as writer:
        result.to_excel(writer, sheet_name="匹配计算结果", index=False)
        method_df.to_excel(writer, sheet_name="方法与阈值说明", index=False)
        summary_df.to_excel(writer, sheet_name="分类汇总", index=False)
        chart_df.to_excel(writer, sheet_name="四象限图数据", index=False)

    wb = load_workbook(output_excel)
    header_fill = PatternFill("solid", fgColor="1F4E78")
    header_font = Font(color="FFFFFF", bold=True)
    thin = Side(style="thin", color="D9E2F3")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                cell.border = border
                cell.alignment = Alignment(vertical="center", wrap_text=True)
                if cell.row == 1:
                    cell.fill = header_fill
                    cell.font = header_font
        for col in ws.columns:
            letter = col[0].column_letter
            max_len = max(len(str(c.value)) if c.value is not None else 0 for c in col)
            ws.column_dimensions[letter].width = min(max(max_len + 2, 10), 36)
        ws.freeze_panes = "A2"

    ws = wb["匹配计算结果"]
    for row in range(2, ws.max_row + 1):
        for col in [2, 3, 6, 8]:
            ws.cell(row=row, column=col).number_format = "0.0000"
    ws.conditional_formatting.add(f"F2:F{ws.max_row}", ColorScaleRule(start_type="min", start_color="F8696B", mid_type="percentile", mid_value=50, mid_color="FFEB84", end_type="max", end_color="63BE7B"))
    ws.conditional_formatting.add(f"H2:H{ws.max_row}", DataBarRule(start_type="min", end_type="max", color="5B9BD5", showValue=True))

    # 四象限图数据页添加散点图
    ws_chart = wb["四象限图数据"]
    chart = ScatterChart()
    chart.title = "老龄化—医疗空间资源四象限图"
    chart.x_axis.title = "U1老龄化综合得分"
    chart.y_axis.title = "U2医疗空间综合得分"
    xvalues = Reference(ws_chart, min_col=2, min_row=2, max_row=ws_chart.max_row)
    yvalues = Reference(ws_chart, min_col=3, min_row=2, max_row=ws_chart.max_row)
    series = Series(yvalues, xvalues, title="省份点位")
    chart.series.append(series)
    chart.width = 18
    chart.height = 10
    ws_chart.add_chart(chart, "H2")

    wb.save(output_excel)


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def write_word(result: pd.DataFrame, output_docx: Path, output_png: Path):
    u1_mean = result["U1老龄化综合得分"].mean()
    u2_mean = result["U2医疗空间综合得分"].mean()
    gap_mean = result["Gap=U2-U1"].mean()
    c_mean = result["耦合度C"].mean()
    gap_counts = result["匹配类型"].value_counts()
    quad_counts = result["四象限类型"].value_counts()
    quad_regions = result.groupby("四象限类型")["地区"].apply(lambda x: "、".join(x)).to_dict()

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)
    doc.styles["Normal"].font.name = "微软雅黑"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    doc.styles["Normal"].font.size = Pt(10.5)

    def add_table(data):
        table = doc.add_table(rows=len(data), cols=len(data[0]))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        for i, row in enumerate(data):
            for j, val in enumerate(row):
                cell = table.rows[i].cells[j]
                cell.text = str(val)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.name = "微软雅黑"
                        r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
                        r.font.size = Pt(9)
                        if i == 0:
                            r.font.bold = True
                            r.font.color.rgb = RGBColor(255, 255, 255)
                if i == 0:
                    set_cell_shading(cell, "1F4E78")
                elif i % 2 == 0:
                    set_cell_shading(cell, "F7FBFF")
        return table

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("老龄化数据与医疗资源空间层面的对应关系分析")
    run.bold = True
    run.font.size = Pt(18)

    doc.add_heading("一、数据来源与分析目标", level=1)
    doc.add_paragraph("本次分析使用上传 Excel 表中已经计算完成的老龄化综合得分 U1 与医疗空间维度综合得分 U2，并按地区进行匹配。")
    doc.add_heading("二、具体步骤与公式", level=1)
    doc.add_paragraph("第一层：匹配差值 Gapᵢ = U₂ᵢ − U₁ᵢ。Gap<0 为资源滞后，Gap≈0 为基本匹配，Gap>0 为空间资源相对富余。")
    doc.add_paragraph("第二层：耦合度 Cᵢ = 2√(U₁ᵢ×U₂ᵢ)/(U₁ᵢ+U₂ᵢ)，用于描述两个系统的同步程度。")
    doc.add_paragraph(f"第三层：四象限图。以 U1 均值 {u1_mean:.4f}、U2 均值 {u2_mean:.4f} 为分界线。")

    doc.add_heading("三、核心计算结果", level=1)
    summary = [
        ["指标", "结果"],
        ["参与匹配地区数", len(result)],
        ["U1老龄化均值", f"{u1_mean:.4f}"],
        ["U2医疗空间均值", f"{u2_mean:.4f}"],
        ["Gap均值", f"{gap_mean:.4f}"],
        ["耦合度C均值", f"{c_mean:.4f}"],
        ["空间资源滞后型地区数", gap_counts.get("空间资源滞后型", 0)],
        ["基本匹配型地区数", gap_counts.get("基本匹配型", 0)],
        ["空间资源相对富余型地区数", gap_counts.get("空间资源相对富余型", 0)],
    ]
    add_table(summary)

    doc.add_heading("四、四象限分类结果", level=1)
    quad_table = [["四象限类型", "地区数量", "地区名单"]]
    for q in ["高老龄—高空间资源", "高老龄—低空间资源", "低老龄—高空间资源", "低老龄—低空间资源"]:
        quad_table.append([q, quad_counts.get(q, 0), quad_regions.get(q, "")])
    add_table(quad_table)
    doc.add_picture(str(output_png), width=Inches(6.4))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("五、结论", level=1)
    doc.add_paragraph(
        f"结果显示，U2 医疗空间维度得分均值 {u2_mean:.4f} 低于 U1 老龄化得分均值 {u1_mean:.4f}，"
        f"Gap 均值为 {gap_mean:.4f}。空间资源滞后型共 {gap_counts.get('空间资源滞后型', 0)} 个地区，"
        "其中高老龄—低空间资源地区应作为空间资源优化和政策倾斜的重点。"
    )
    doc.save(output_docx)


def main():
    if not INPUT_FILE.exists():
        raise FileNotFoundError(f"未找到输入文件：{INPUT_FILE}")

    base_df = read_score_data(INPUT_FILE)
    result = add_model_results(base_df)
    make_quadrant_plot(result, OUTPUT_PNG)
    write_excel(result, OUTPUT_EXCEL)
    write_word(result, OUTPUT_DOCX, OUTPUT_PNG)

    print("分析完成，输出文件如下：")
    print(OUTPUT_EXCEL)
    print(OUTPUT_PNG)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
